import os
from os import environ as env
import click
from docx import Document
from docx.table import Table, _Cell, _Column
from docx.styles.styles import Styles
from docx.document import Document as _Document
from docx.shared import Pt
from pathlib import Path
from pandas import DataFrame

from src.utils.markdown import markdown_to_word, parse_md
from src.utils.math import add_tuples


os.environ["ROOT_DIR"] = str(Path(__file__).parent.parent.resolve())

# Absolute Path of course content folder from env
assert "COURSE_CONTENT" in env, "COURSE_CONTENT is undefined"
assert "OUTPUT_LOCATION" in env, "OUTPUT_LOCATION is undefined"

COURSE_CONTENT = Path(env["COURSE_CONTENT"]).resolve()
OUTPUT_LOCATION = Path(env["OUTPUT_LOCATION"]).resolve()

# Source code locations:
ROOT = env["ROOT_DIR"]  # repo root location
TEMPLATES = Path("templates/")

# Implementation Specific
TEMPLATE = TEMPLATES / Path("Assessment Task Tool (F122A12).docx")
OUTPUT_FILE = Path("Assessment Task Tool (F122A12).docx")

# Relative Path of Content Files (Input and Output):
ASSESSMENTS = Path("2 KAD/5 Assess Tool/")


import re
from typing import List, Dict


def parse_markdown_headers(md_content: str) -> List[Dict[str, str]]:
    """
    Parses a Markdown string into an iterable of sections based on Markdown headers.

    :param md_content: A string containing the Markdown content.
    :return: A list of dictionaries with 'header' and 'content' keys.
    """

    # Define a regex to match markdown headers
    header_regex = re.compile(r"^(#{1})\s+(.*)", re.MULTILINE)

    sections = []
    last_pos = 0
    for match in header_regex.finditer(md_content):
        # Extract header level and text
        header_level = len(match.group(1))
        header_text = match.group(2).strip()

        # Find the position of the header
        start_pos = match.start()
        # Get content up to this header
        content = md_content[last_pos:start_pos].strip()

        # If there is a previous section, update its content
        if sections:
            sections[-1]["content"] = content

        # Create a new section for the current header
        section = {"header": header_text, "content": "", "level": header_level}
        sections.append(section)

        last_pos = match.end()

    # Add the content for the last section
    if sections:
        sections[-1]["content"] = md_content[last_pos:].strip()

    return sections


def assess_tool(course_directory: Path, output_location: Path):
    assert course_directory.is_dir()
    assert output_location.is_dir()

    doc: _Document = Document(ROOT / TEMPLATE)
    styles: Styles = doc.styles

    assessments = course_directory / ASSESSMENTS
    for assessment in assessments.rglob("assessment.md"):
        output: Path = (
            output_location / ASSESSMENTS / Path(assessment.parent.name) / OUTPUT_FILE
        )

        markdown = parse_md(assessment)

        sections = parse_markdown_headers(markdown.content)

        for idx, section in enumerate(sections):
            table_number = 1 + idx
            table: Table = doc.tables[table_number - 1]

            cell: _Cell = table.cell(0, 0)
            cell.text = ""
            markdown_to_word(section.get("content", ""), doc, cell)

        for checklist in markdown.get("observation_checklist", []) or []:
            doc.add_page_break()
            doc.add_heading("Observation Checklist", 2)
            table = doc.add_table(0, 0)
            table.autofit = True
            row = table.add_row()
            for column_idx, column in enumerate(checklist.keys()) or []:
                _column: _Column = table.add_column(20)
                table.cell(*(0, column_idx)).text = column

                rows = checklist.get(column) or []

                for row_idx, value in enumerate(rows):
                    if table.rows is None:
                        pass
                    if value is None:
                        value = ""
                    if row_idx + 1 >= len(table.rows):
                        table.add_row()
                    table.cell(*(row_idx + 1, column_idx)).text = value

        output.parent.mkdir(exist_ok=True, parents=True)
        doc.save(output)


@click.command()
# @click.argument("course_directory", type=click.Path(exists=True, path_type=Path))
def run_cli():
    """
    CLI tool to write YAML header data from Markdown file to Word document as custom properties.
    """
    assess_tool(COURSE_CONTENT, OUTPUT_LOCATION)


if __name__ == "__main__":
    run_cli()
