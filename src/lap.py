import os
from os import environ as env
import click
from docx import Document
from docx.table import Table, _Cell
from docx.styles.styles import Styles
from docx.shared import Pt
from pathlib import Path
from pandas import DataFrame

from src.utils.markdown import markdown_to_word, parse_md
from src.utils.math import add_tuples


os.environ["ROOT_DIR"] = str(Path(__file__).parent.parent.resolve())

# Absolute Path of course content folder from env
assert "COURSE_CONTENT" in env, "COURSE_CONTENT is undefined"
assert "OUTPUT_LOCATION" in env, "OUTPUT_LOCATION is undefined"
COURSE_CONTENT = Path(env["COURSE_CONTENT"]).resolve()
OUTPUT_LOCATION = Path(env["OUTPUT_LOCATION"]).resolve()

# Source code locations:
ROOT = env["ROOT_DIR"] # repo root location
TEMPLATES = Path("templates/")

# Implementation Specific
TEMPLATE = TEMPLATES / Path("Learning and Assessment Plan (F122A14).docx")
OUTPUT_FILE = Path("2 KAD/1 LAP/Learning and Assessment Plan (F122A14).docx")

# Relative Path of Content Files:
TOPICS = Path("2 KAD/1 LAP/topics.md")
FIELDS = Path("2 KAD/1 LAP/fields.md")
RESOURCES = Path("2 KAD/1 LAP/resources.md")
ELEMENTS = Path("2 KAD/1 LAP/elements.md")


import re
from typing import List, Dict




def parse_markdown_headers(md_content:str) -> List[Dict[str, str]]:
    """
    Parses a Markdown string into an iterable of sections based on Markdown headers.

    :param md_content: A string containing the Markdown content.
    :return: A list of dictionaries with 'header' and 'content' keys.
    """

    # Define a regex to match markdown headers
    header_regex = re.compile(r"^(#{1,6})\s+(.*)", re.MULTILINE)

    sections = []
    last_pos = 0
    for match in header_regex.finditer(md_content):
        # Extract header level and text
        header_level = len(match.group(1))
        header_text = match.group(2).strip()

        # Find the position of the header
        start_pos = match.start()
        # Get content up to this header
        content = md_content[last_pos:start_pos].strip()

        # If there is a previous section, update its content
        if sections:
            sections[-1]["content"] = content

        # Create a new section for the current header
        section = {"header": header_text, "content": "", "level": header_level}
        sections.append(section)

        last_pos = match.end()

    # Add the content for the last section
    if sections:
        sections[-1]["content"] = md_content[last_pos:].strip()

    return sections



def lap(course_directory: Path, output_location: Path):
    assert course_directory.is_dir()
    assert output_location.is_dir()
    
    doc = Document(ROOT / TEMPLATE)
    styles: Styles = doc.styles

    output_location.mkdir(parents=True, exist_ok=True)

    # Populate Fields
    fields = parse_md(course_directory / FIELDS)
    
    # Table 1
    ## Qualification
    ## Delivery Period
    ## Cluster Name
    table_number = 1 
    table: Table = doc.tables[table_number - 1]
    
    table.cell(*(0,1)).text = fields.get("qualification_national_code_and_title","")
    table.cell(*(1,1)).text = fields.get("delivery_period","")
    table.cell(*(2,1)).text = fields.get("cluster_name","")

    # Table 2
    # National ID
    # Name of Unit
    table_number = 2 
    table: Table = doc.tables[table_number - 1]
    
    for unit_number, unit in enumerate(fields.get("units","")):
        table.cell(*(1 + unit_number, 1)).text = unit.get("name","")
        table.cell(*(1 + unit_number, 0)).text = unit.get("id","")
    
    
    table.cell(*(5,1)).text = fields.get("delivery_location/s","")
    
    # Table 3
    table_number = 3 
    table: Table = doc.tables[table_number - 1]
    table.cell(*(1,0)).paragraphs[0].add_run(fields.get("student_to_supply",""))
    table.cell(*(2,0)).paragraphs[0].add_run(fields.get("college_to_supply",""))


    for lecturer_number, lecturer in enumerate(fields.get("lecturers","")):
        if lecturer_number > 1:
            table.add_row()
            
        table.cell(*(4+lecturer_number,0)).text = lecturer.get("name","")
        table.cell(*(4+lecturer_number,1)).text = lecturer.get("phone","")
        table.cell(*(4+lecturer_number,2)).text = lecturer.get("email","")
        table.cell(*(4+lecturer_number,3)).text = lecturer.get("contact_time","")
        table.cell(*(4+lecturer_number,4)).text = lecturer.get("campus/room","")
        
        
    
    # Table 4
    table_number = 4 
    table: Table = doc.tables[table_number - 1]
    assessment_number = 0
    for assessment in fields.get("assessments",""):
        if assessment_number > 3:
            table.add_row()
        row = 1 + assessment_number
        table.cell(*(row, 0)).text = f"Assessment {assessment_number + 1}"
        table.cell(*(row, 1)).text = assessment.get("title","")
        table.cell(*(row, 2)).text = assessment.get("due_date","")
        
        assessment_number += 1
        
    # Table 5
    table_number = 5 
    table: Table = doc.tables[table_number - 1]
    # Currently I have no Idea how to change the state of the checkboxes used in the lap
    # it is probably easier to have a different template for each kind of lap.
    # cell = table.cell(*(2,0))
    # print(dir(cell))
    # unchecked_checkbox_character = u'\u2610'
    # checked_checkbox_character = u'\u2611'
    # for p in cell.paragraphs:
    #     p.add_run(unchecked_checkbox_character)
    #     print(checked_checkbox_character in p.text)
    
    
    # Table 6
    # Session Topics
    table_number = 6
    table: Table = doc.tables[table_number - 1]
    
    parsed_md = parse_md(course_directory / TOPICS)
    elements = parse_md(course_directory / ELEMENTS)
    resources = parse_md(course_directory / RESOURCES).content.split("---")
    
    topics = parse_markdown_headers(parsed_md.content)
    hours_coords = (2, 1)
    element_coords = (2, 2)
    topic_coords = (2, 3)
    resources_coords = (2, 4)
    ooch_coords = (2, 6)
    # table.autofit = True
    for idx, topic in enumerate(topics):
        POINTER = (idx, 0)
        # Populate Topics
        coords = add_tuples(POINTER, topic_coords)
        cell: _Cell = table.cell(*coords)
        cell.text = ""
        cell.paragraphs[-1].text = topic.get("header") 
        cell.paragraphs[-1].style = styles[f"Heading {topic.get("level", 1)}"] 
        markdown_to_word(topic.get("content"), doc, cell)
        
        # Populate Session Hours
        coords = add_tuples(POINTER, hours_coords)
        cell: _Cell = table.cell(*coords)
        cell.paragraphs[-1].text = str(parsed_md.get("session_hours", 0))
        
        # Populate Out of class hours
        coords = add_tuples(POINTER, ooch_coords)
        cell: _Cell = table.cell(*coords)
        cell.paragraphs[-1].text = str(parsed_md.get("out_of_class_hours",0))

        # Populate Knowledge Evidence
        font_name = "Arial"
        font_size = Pt(8)
        coords = add_tuples(POINTER, element_coords)
        cell: _Cell = table.cell(*coords)
        
        sessions:list = elements.get("sessions", [])
        try:
            if any([len(knowledge) > 0 for unit in sessions[idx] for knowledge in unit.get("knowledge",[]) or []]):
                run = cell.paragraphs[-1].add_run("Knowledge Element")
                run.bold = True
                run.font.name = font_name
                run.font.size = font_size
                for unit in sessions[idx]:
                    knowledge = unit.get("knowledge", []) or []
                    if len(knowledge) > 0:
                        p = cell.add_paragraph(unit.get("name") + ":")
                        run = p.runs[-1]
                        run.bold = True
                        run.font.name = font_name
                        run.font.size = font_size
                        for element in knowledge: 
                            run = p.add_run(f" {element} ")
                            run.font.name = font_name
                            run.font.size = font_size
        except Exception as e:
            print(e)
            
        # Learning Resources
        # for resource in resources:
        coords = add_tuples(POINTER, resources_coords)
        cell: _Cell = table.cell(*coords)
        markdown_to_word(resources[idx], doc, cell)

        
        table.cell(*(22, 1)).text = str(parsed_md.get("total_session_hours"))
        table.cell(*(22, 6)).text = str(parsed_md.get("total_out_of_class_hours"))
        table.cell(*(23, 5)).text = str(parsed_md.get("total_training"))
        
        # table.cell(*coords).add_paragraph(topic.get("content"), styles[f"Normal"])

    # for row in doc.tables[5].rows:
    #     for cell in row.cells:
    #         cell.text = "1"
    (output_location/OUTPUT_FILE).parent.mkdir(exist_ok=True, parents=True)
    doc.save(output_location/OUTPUT_FILE)


@click.command()
# @click.argument("course_directory", type=click.Path(exists=True, path_type=Path))
def run_cli():
    """
    CLI tool to write YAML header data from Markdown file to Word document as custom properties.
    """
    lap(COURSE_CONTENT, OUTPUT_LOCATION)


if __name__ == "__main__":
    run_cli()
