from itertools import chain
import os
from os import environ as env
import click
from docx import Document
from docx.shared import Pt, Inches
from docx.table import Table, _Cell, _Column
from docx.styles.styles import Styles
from docx.styles.style import _ParagraphStyle
from docx.enum.style import WD_STYLE_TYPE, WD_BUILTIN_STYLE as WD_STYLE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.document import Document as _Document
from docx.shared import Pt
from docx.section import _Header, _Footer, Section, Sections
from docx.text.paragraph import Paragraph
from pathlib import Path
from pandas import DataFrame

from src.utils.markdown import markdown_to_word, parse_md
from src.utils.math import add_tuples
from src.utils.uoc import UnitOfCompetency
from docx.enum.text import WD_ALIGN_PARAGRAPH
from frontmatter import Post


# normal_bold = _ParagraphStyle()
# normal_bold.font.bold = True

os.environ["ROOT_DIR"] = str(Path(__file__).parent.parent.resolve())

# Absolute Path of course content folder from env
assert "COURSE_CONTENT" in env, "COURSE_CONTENT is undefined"
assert "OUTPUT_LOCATION" in env, "OUTPUT_LOCATION is undefined"

COURSE_CONTENT = Path(env["COURSE_CONTENT"]).resolve()
OUTPUT_LOCATION = Path(env["OUTPUT_LOCATION"]).resolve()

# Source code locations:
ROOT = env["ROOT_DIR"]  # repo root location
TEMPLATES = Path("templates/")

# Implementation Specific
TEMPLATE = TEMPLATES / Path("Assessment Mapping Matrix (F122A8).docx")
OUTPUT_FILE = Path("Assessment Mapping Matrix (F122A8).docx")

# Relative Path of Content Files (Input and Output):
ASSESSMENTS = Path("2 KAD/5 Assess Tool/")
MAPPING_MATRIX = Path("2 KAD/7 Assess Mapping Matrix/")


import re
from typing import List, Dict


def parse_markdown_headers(md_content: str) -> List[Dict[str, str]]:
    """
    Parses a Markdown string into an iterable of sections based on Markdown headers.

    :param md_content: A string containing the Markdown content.
    :return: A list of dictionaries with 'header' and 'content' keys.
    """

    # Define a regex to match markdown headers
    header_regex = re.compile(r"^(#{1})\s+(.*)", re.MULTILINE)

    sections = []
    last_pos = 0
    for match in header_regex.finditer(md_content):
        # Extract header level and text
        header_level = len(match.group(1))
        header_text = match.group(2).strip()

        # Find the position of the header
        start_pos = match.start()
        # Get content up to this header
        content = md_content[last_pos:start_pos].strip()

        # If there is a previous section, update its content
        if sections:
            sections[-1]["content"] = content

        # Create a new section for the current header
        section = {"header": header_text, "content": "", "level": header_level}
        sections.append(section)

        last_pos = match.end()

    # Add the content for the last section
    if sections:
        sections[-1]["content"] = md_content[last_pos:].strip()

    return sections


def mapping_matrix(course_directory: Path, output_location: Path):
    assert course_directory.is_dir()
    assert output_location.is_dir()

    assessments = course_directory / ASSESSMENTS
    unit_assessment_mapping = {}

    for assessment in sorted(assessments.rglob("assessment.md")):
        if not assessment.is_file():
            continue

        markdown = parse_md(assessment)

        name = markdown.get("name")
        units = markdown.get("units")

        for unit in units:
            ## Initialize unit mapping matrix if needed:
            unit_assessment_mapping.setdefault(
                unit["id"],
                {
                    "assessments": [],
                    "unit": unit,
                    "qualification": markdown.get(
                        "qualification_national_code_and_title"
                    ),
                },
            )
            ## add assessment to unit mapping matrix
            unit_assessment_mapping.get(unit["id"]).get("assessments").append(markdown)

    for unit_index, (id, mapping_matrix) in enumerate(unit_assessment_mapping.items()):
        doc: _Document = Document(ROOT / TEMPLATE)
        styles: Styles = doc.styles

        ## Header Formatting

        unit = mapping_matrix.get("unit")
        header: _Header = doc.sections[0].header
        table_header: Table = header.tables[0]

        # Set Unit national codes and titles
        cell: _Cell = table_header.cell(1, 1)
        cell.text = f'{unit.get("id")} {unit.get("name")}'

        # Set qualification national codes and titles
        cell: _Cell = table_header.cell(0, 1)
        cell.text = f'{mapping_matrix.get("qualification")}'

        ## Mapping Matrix
        table = doc.tables[0]
        uoc: UnitOfCompetency = UnitOfCompetency(id)

        # Elements
        elements: dict = uoc.data.elements_and_criteria
        for element_index, (element, criteria) in enumerate(elements.items()):
            element_header = next(
                (
                    index
                    for index, cell in enumerate(table.column_cells(0))
                    if f"Element {element_index + 1}" in cell.text
                )
            )
            sections = criteria.strip().split("\n")
            for index, criterium in enumerate(sections):
                table.cell(element_header + 1 + index, 0).text = criterium

        # Knowledge
        knowledge_elements = uoc.parse_knowledge_criteria()
        knowledge_header = next(
            (
                index
                for index, cell in enumerate(table.column_cells(0))
                if f"Required Knowledge or Knowledge Evidence" in cell.text
            )
        )
        for index, element in enumerate(knowledge_elements.keys()):
            cell = table.cell(knowledge_header + 1 + index, 0)
            cell.text = element
            cell.paragraphs[0].runs[0].bold = True
            if len(knowledge_elements[element]) > 0:
                for sub_element in knowledge_elements[element]:
                    paragraph = cell.add_paragraph()
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.add_run(sub_element)

        # Performance Evidence
        performance = uoc.parse_performance_evidence()
        performance_header = next(
            (
                index
                for index, cell in enumerate(table.column_cells(0))
                if f"Required Skills or Performance Evidence" in cell.text
            )
        )
        for index, element in enumerate(performance.keys()):
            cell = table.cell(performance_header + 1 + index, 0)
            cell.text = element
            if ":" in cell.text:
                cell.paragraphs[0].runs[0].bold = True
            if len(performance[element]) > 0:
                for sub_element in performance[element]:
                    paragraph = cell.add_paragraph()
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.add_run(sub_element)

        # Assessment Conditions
        assessment_conditions = uoc.parse_assessment_conditions()
        ac_header = next(
            (
                index
                for index, cell in enumerate(table.column_cells(0))
                if f"Assessment Conditions" in cell.text
            )
        )
        rows = []
        for index, element in enumerate(assessment_conditions.keys()):
            row_index = ac_header + 1 + index
            rows.append(row_index)
            cell = table.cell(row_index, 0)
            cell.text = element
            if ":" in cell.text:
                cell.paragraphs[0].runs[0].bold = True
            if len(assessment_conditions[element]) > 0:
                for sub_element in assessment_conditions[element]:
                    paragraph = cell.add_paragraph()
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.add_run(sub_element)

            cell = table.cell(row_index, 1)
            for other_cell in (
                table.cell(row_index, index) for index in range(2, len(table.columns))
            ):
                cell.merge(other_cell)

        rows.reverse()
        first_row = rows.pop()
        rows.reverse()
        for other_cell in (table.cell(row, 1) for row in rows):
            table.cell(first_row, 1).merge(other_cell)

        for other_cell in (table.cell(row, 0) for row in rows):
            table.cell(first_row, 0).merge(other_cell)

        table.cell(first_row, 1).text = (
            "\n".join(assessment_conditions.keys())
            .replace("must be", "are")
            .replace("must", "always")
        )

        assessments: list[Post] = mapping_matrix.get("assessments")
        for assessment_index, assessment in enumerate(assessments):
            cell: _Cell = table.cell(0, assessment_index + 1)
            paragraph: Paragraph = cell.paragraphs[0]
            paragraph.clear()
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Set up columns
            paragraph.text = f"Assessment Task {assessment_index + 1}"
            paragraph.style = doc.styles["Heading 3"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set up Assessment Title
            table.cell(1, assessment_index + 1).text = assessment.get("name")

            # Mapping
            mapping: list = assessment.get("mapping", []) or []
            elements
            knowledge_elements
            ### Preprocess mapping into element -> question numbers

            ## TODO: Update
            ## Set Element mapping
            for element_index, (element, criteria) in enumerate(elements.items()):
                element_header = next(
                    (
                        index
                        for index, cell in enumerate(table.column_cells(0))
                        if f"Element {element_index + 1}" in cell.text
                    )
                )
                sections = criteria.strip().split("\n")
                for index, criterium in enumerate(sections):
                    key: float = float(criterium[:3])
                    question_mapping: str = ", ".join(
                        (
                            str(question_index + 1)
                            for question_index, question in enumerate(mapping)
                            if key
                            in (((question or {}).get("criteria") or {}).get(id) or [])
                        )
                    )
                    table.cell(
                        element_header + 1 + index, 1 + assessment_index
                    ).text = question_mapping

            ## Set Knowledge mapping
            knowledge_header
            for knowledge_index, element in enumerate(knowledge_elements.keys()):
                cell = table.cell(
                    knowledge_header + 1 + knowledge_index, 1 + assessment_index
                )
                key: int = int(knowledge_index + 1)
                question_mapping: str = ", ".join(
                    (
                        str(question_index + 1)
                        for question_index, question in enumerate(mapping)
                        if key
                        in (((question or {}).get("knowledge") or {}).get(id) or [])
                    )
                )

                cell.text = question_mapping
                # cell.paragraphs[0].runs[0].bold = True

            ## Set Performance & Skills Mapping
            ## TODO: THIS DOESN'T WORK!!! FiX IT PLEASE!!!
            performance_header = next(
                (
                    index
                    for index, cell in enumerate(table.column_cells(0))
                    if f"including evidence of the ability to:" in cell.text
                )
            )

            skills_header = next(
                (
                    index
                    for index, cell in enumerate(table.column_cells(0))
                    if f"In the course of the above, the candidate must:" in cell.text
                )
            )

            for performance_index, element_number in enumerate(
                chain(
                    *[
                        ((question or {}).get("performance") or {}).get(id) or []
                        for question in mapping
                    ]
                )
            ):
                cell = table.cell(
                    performance_header + 1 + performance_index, 1 + assessment_index
                )
                key: int = int(performance_index + 1)
                question_mapping: str = ", ".join(
                    (
                        str(question_index + 1)
                        for question_index, question in enumerate(mapping)
                        if key
                        in (((question or {}).get("performance") or {}).get(id) or [])
                    )
                )
                cell.text = question_mapping

            for skills_index, element_number in enumerate(
                chain(
                    *[
                        ((question or {}).get("skills") or {}).get(id) or []
                        for question in mapping
                    ]
                )
            ):
                cell = table.cell(
                    skills_header + 1 + skills_index, 1 + assessment_index
                )
                key: int = int(skills_index + 1)
                question_mapping: str = ", ".join(
                    (
                        str(question_index + 1)
                        for question_index, question in enumerate(mapping)
                        if key in (((question or {}).get("skills") or {}).get(id) or [])
                    )
                )
                cell.text = question_mapping

        # TODO: auto insert uoc elements etc

        # TODO: set each column up for every assessment (names etc..)

        # TODO: Implement actual mapping of assessment elements to criteria etc

        # TODO: Implement main disclosure statements
        # table.autofit = True
        output: Path = output_location / MAPPING_MATRIX / (id + " " + str(OUTPUT_FILE))
        output.parent.mkdir(exist_ok=True, parents=True)
        doc.save(output)


@click.command()
# @click.argument("course_directory", type=click.Path(exists=True, path_type=Path))
def run_cli():
    """
    CLI tool to write YAML header data from Markdown file to Word document as custom properties.
    """
    mapping_matrix(COURSE_CONTENT, OUTPUT_LOCATION)


if __name__ == "__main__":
    run_cli()
