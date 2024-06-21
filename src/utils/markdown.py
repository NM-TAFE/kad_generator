from pathlib import Path
import re
from docx.shared import Pt, Inches
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.text.paragraph import Paragraph
import frontmatter

## General Markdown functions


def parse_md(path: Path):
    assert path.is_file()
    assert path.exists()
    # Load the markdown file and parse the front matter
    with open(path, "r", encoding="utf-8") as file:
        parsed_md = frontmatter.load(file)
        parsed_md.content = re.sub(r"<!--.*-->", "", parsed_md.content)
    return parsed_md


## Markdown to Word Style Mapping:
## ! Warning: Avoid defining overlapping regex, this is not supported by current implementation (afaik)
## ! It will lead to duplicated text.
## Parsing markdown effectively is quite complex thus we have only implemented some of the basic syntax here
## Custom style mappings can be added as needed
MARKDOWN_STYLES = {
    "h1": {"regex": re.compile(r"^#{1} (.*)", re.MULTILINE), "style": "Heading 1"},
    "h2": {"regex": re.compile(r"^#{2} (.*)", re.MULTILINE), "style": "Heading 2"},
    "h3": {"regex": re.compile(r"^#{3} (.*)", re.MULTILINE), "style": "Heading 3"},
    "h4": {"regex": re.compile(r"^#{4} (.*)", re.MULTILINE), "style": "Heading 4"},
    "h5": {"regex": re.compile(r"^#{5} (.*)", re.MULTILINE), "style": "Heading 5"},
    "h6": {"regex": re.compile(r"^#{6} (.*)", re.MULTILINE), "style": "Heading 6"},
    "bold/italic": {
        "regex": re.compile(r"(\*{1,2})([^*]*?)(\*{1,2})"),
        "style": "bold/italic",
    },
    "code": {"regex": re.compile(r"`{3}([^`]*)`{3}"), "style": "code"},
    "bullets": {
        "regex": re.compile(r"^([' ',\t]*)[*\-+]\s(.*)$", re.MULTILINE),
        "style": "List Bullet",
    },
    # "numbers": {
    #     "regex": re.compile(r"^([' ',\t]*)\d{1,}\.\s(.*)$", re.MULTILINE),
    #     "style": "List Number",
    # },
    "link": {"regex": re.compile(r"(?<!!)\[(.*)\]\((.*)\)")},
    "image": {"regex": re.compile(r"!\[(.*)\]\((.*)\)")},
    "linebreak": {"regex": re.compile(r"^\-{3}$", re.MULTILINE)},
    # Add more patterns if needed, like lists, links, etc.
}


def apply_markdown_style(document, text, parent=None):
    """
    Apply Markdown styles to text within a given parent or a new paragraph in the document.

    :param document: docx Document object.
    :param text: Text string containing Markdown content.
    :param parent: Parent container for the text runs (like a table cell or paragraph).
    :return: The paragraph to which the styles were applied.
    """
    if parent is None:
        parent = document

    def add_paragraph(style=None) -> Paragraph:
        if parent is None:
            return document.add_paragraph(style=style)
        else:
            return parent.add_paragraph(style=style)

    def print_matches(text, sorted_matches):
        # Variables to track lists and their levels
        bullet_level = ""

        # If no parent is given, append a new paragraph to the document.
        paragraph = document.paragraphs[-1] if parent is None else parent.paragraphs[-1]

        # Keep track of the last match position.
        last_idx = 0
        last_style = None

        for match, style, wstyle in sorted_matches:
            # match: Regex match
            # style: markdown style id
            # wstyle: Correlated Word style (if exists) Must be Paragraph style

            remaining: str = text[last_idx : match.start()]

            paragraph.add_run(remaining)

            last_idx = match.start()

            if style == "bullets":

                indentation = len(match.group(1).replace("\t", "  ")) // 2

                paragraph = add_paragraph(
                    style=f'{wstyle}{f" {indentation + 1}" if indentation > 0 else ""}'
                )
                paragraph.add_run(match.group(2))
                last_idx = match.end()
                last_style = style

            else:
                # Add the text before the markdown pattern.
                paragraph.add_run(text[last_idx : match.start()])

                # Add the styled text for the matched markdown pattern.
                if wstyle == "bold/italic":
                    paragraph = add_paragraph()
                    styled_run = paragraph.add_run(match.group(2))
                    styled_run.bold = True if len(match.group(1)) == 2 else False
                    styled_run.italic = True if len(match.group(1)) == 1 else False
                elif wstyle == "code":
                    styled_run = paragraph.add_run(match.group(1))
                    styled_run.font.name = "Arial"
                    styled_run.font.size = Pt(10)
                elif style == "link":
                    add_hyperlink(paragraph, match.group(1), match.group(2))
                elif style == "image":
                    try:
                        paragraph.add_run().add_picture(match.group(2))
                    except:
                        paragraph.add_run(match.group(1))
                elif style == "linebreak":
                    # We process line breaks as 10 empty paragraphs instead
                    for _ in range(10):
                        paragraph = add_paragraph()
                else:
                    paragraph = add_paragraph(style=document.styles[wstyle])
                    styled_run = paragraph.add_run(match.group(1))

                last_idx = match.end()
                last_style = style
        else:
            remaining: str = text[last_idx:]

            # paragraph = add_paragraph()
            paragraph.add_run(remaining)

            return paragraph

    for line in text.split("\n"):
        matches = [
            (match, style, pattern_info.get("style", style))
            for style, pattern_info in MARKDOWN_STYLES.items()
            for match in pattern_info["regex"].finditer(line)
        ]
        sorted_matches = sorted(matches, key=lambda x: x[0].start())
        if len(sorted_matches) == 1:
            print_matches(line, sorted_matches)
        elif len(sorted_matches) == 0:
            paragraph = add_paragraph()
            print_matches(line, sorted_matches)
        else:
            # We in trouble
            raise NotImplementedError(
                "Multiple styles in a single line not yet implemented.", sorted_matches
            )


def markdown_to_word(doc_content, document, parent=None):
    """
    Parse the given Markdown content and apply styles to a Word document or a specified parent container.

    :param doc_content: String containing Markdown content.
    :param document: docx Document object.
    :param parent: Optional. Parent container such as a table cell in the document.
                   If none is provided, new paragraphs are added to the document.
    """
    # Split content into Markdown blocks.
    # blocks = doc_content.split("\n")

    # Function to add a paragraph either to the document root or within a specified parent.
    def add_paragraph(style=None):
        if parent is None:
            return document.add_paragraph(style=style)
        else:
            return parent.add_paragraph(style=style)

        # for block in blocks:
        #     # Check for Markdown headers and apply accordingly.
        #     for header_pattern in ("h1", "h2", "h3", "h4", "h5"):
        #         header = MARKDOWN_STYLES[header_pattern]
        #         match = header["regex"].match(block.lstrip())
        #         if match:
        #             paragraph = add_paragraph(style=document.styles[header["style"]])
        #             paragraph.add_run(match.group(1))
        #             paragraph = add_paragraph()
        #             break
        #     else:  # If not a header, apply Markdown styles to a new or existing parent.

    apply_markdown_style(document, doc_content, parent=parent)
    # paragraph = add_paragraph()


def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The Paragraph object where the hyperlink will be added.
    :param text: The text displayed for the hyperlink.
    :param url: The destination URL for the hyperlink.
    :return: The hyperlink object.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        qn("r:id"),
        r_id,
    )

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set the style for the hyperlink (typically blue underlined text)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000EE")  # Color value.
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink
